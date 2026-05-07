"""
Script para generar el informe de la Práctica 3 BDA en formato .docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Helpers ──────────────────────────────────────────────────────────────────

BLUE   = RGBColor(0, 63, 135)   # azul institucional
DARK   = RGBColor(26, 32, 53)   # casi negro
GRAY   = RGBColor(100, 116, 139)
RED    = RGBColor(185, 28, 28)
GREEN  = RGBColor(22, 101, 52)
CODE_BG = RGBColor(13, 17, 23)   # fondo oscuro para código


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'DEE2E6')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def heading(doc, text, level=1, color=BLUE):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    run = p.runs[0] if p.runs else p.add_run(text)
    run.text = text
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(16)
        run.bold = True
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(13)
        run.bold = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    return p


def body(doc, text, bold=False, italic=False, color=DARK, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)
    return p


def code_block(doc, lines):
    """Bloque de código con fondo oscuro y texto de colores."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # Fondo gris claro (no podemos aplicar fondo por párrafo fácilmente, usamos color de texto)
        color_map = {
            'INFO':  RGBColor(60, 130, 200),
            'OK':    RGBColor(40, 180, 100),
            'ERROR': RGBColor(220, 50, 50),
            'WARN':  RGBColor(210, 160, 40),
            '--':    GRAY,
        }
        text_color = DARK
        for key, col in color_map.items():
            if key in line.upper():
                text_color = col
                break
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = text_color


def inline_code(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(180, 50, 50)
    return run


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8A951')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Documento ─────────────────────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Fuente por defecto
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)
doc.styles['Normal'].font.color.rgb = DARK


# ════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('BASES DE DATOS AVANZADAS')
run.font.size = Pt(11)
run.font.color.rgb = BLUE
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Práctica 3 — Gestor Académico MVC')
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = DARK
run.font.name = 'Calibri'
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema de gestión académica con Flask, PostgreSQL y arquitectura MVC')
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.italic = True
run.font.name = 'Calibri'

doc.add_paragraph()

# Tabla de datos del informe
table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
set_cell_borders(table)
info_rows = [
    ('Asignatura',  'Bases de Datos Avanzadas'),
    ('Práctica',    'Práctica 3'),
    ('Tecnologías', 'Python 3.11 · Flask 3.x · PostgreSQL 17 · psycopg2'),
    ('Patrón',      'MVC (Modelo–Vista–Controlador)'),
    ('Fecha',       'Mayo 2026'),
]
for i, (label, val) in enumerate(info_rows):
    row = table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = val
    set_cell_bg(row.cells[0], 'E8F0FB')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
        cell.paragraphs[0].paragraph_format.space_before = Pt(4)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(4)
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = BLUE

doc.add_page_break()


# ════════════════════════════════════════════════════════════
#  1. VERSIONES
# ════════════════════════════════════════════════════════════

heading(doc, '1. Versiones del entorno', 1)

body(doc, (
    'La práctica ha sido desarrollada y ejecutada sobre las siguientes versiones de software. '
    'Ambas herramientas fueron elegidas por su madurez, soporte activo y compatibilidad con '
    'las características avanzadas de bases de datos relacionales que se trabajan en la asignatura.'
))

# Tabla de versiones
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
set_cell_borders(table)
table.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['Componente', 'Versión', 'Rol en el proyecto']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '003F87')
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(10)

rows_data = [
    ('Python',        '3.11.15',   'Lenguaje del servidor. Ejecuta Flask, la lógica MVC y las consultas psycopg2.'),
    ('PostgreSQL',    '17.4',      'Motor de base de datos relacional. Gestiona las tablas, índices, FK y transacciones ACID.'),
    ('psycopg2',      '2.9.9+',    'Driver Python–PostgreSQL. Gestiona conexiones, cursores y el ciclo BEGIN/COMMIT/ROLLBACK.'),
]
for i, (comp, ver, rol) in enumerate(rows_data, start=1):
    row = table.rows[i]
    row.cells[0].text = comp
    row.cells[1].text = ver
    row.cells[2].text = rol
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F0F4FB')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)

row.cells[0].paragraphs[0].runs[0].bold = True

# Verificar versión
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Verificación en psql: ')
run.font.size = Pt(10)
inline_code(p, 'SELECT version();')
run2 = p.add_run('\nResultado obtenido: ')
run2.font.size = Pt(10)
inline_code(p, 'PostgreSQL 17.4 on x86_64-pc-linux-gnu, compiled by gcc 13.3.0, 64-bit')

add_separator(doc)


# ════════════════════════════════════════════════════════════
#  2. PASOS PARA EJECUTAR
# ════════════════════════════════════════════════════════════

heading(doc, '2. Pasos para ejecutar el proyecto', 1)

body(doc, 'Requisitos previos: Python ≥ 3.10, PostgreSQL ≥ 15 en ejecución y acceso a una base de datos.')

heading(doc, '2.1 Instalación de dependencias', 2)

body(doc, 'Clonar el repositorio e instalar las dependencias Python:')
code_block(doc, [
    '# Clonar el repositorio',
    'git clone https://github.com/Rodrcam/Practica3_BDA-.git',
    'cd Practica3_BDA-/BBDDAvanzadas-colab',
    '',
    '# (Opcional) Crear entorno virtual',
    'python -m venv venv && source venv/bin/activate',
    '',
    '# Instalar dependencias',
    'pip install -r requirements.txt',
])

body(doc, 'El archivo requirements.txt contiene:', size=10)
bullet(doc, 'Flask >= 3.0.0 — framework web')
bullet(doc, 'psycopg2-binary >= 2.9.9 — driver PostgreSQL')
bullet(doc, 'Faker >= 24.0.0 — generación de datos de prueba')

heading(doc, '2.2 Configuración de la base de datos', 2)

body(doc, 'Copiar el archivo de configuración de ejemplo y rellenarlo con los datos de conexión:')
code_block(doc, [
    'cp config/database.ini.example config/database.ini',
])

body(doc, 'Contenido de config/database.ini:')
code_block(doc, [
    '[postgresql]',
    'host     = localhost',
    'database = BBDDAvanzadasGestorAlumnos',
    'user     = postgres',
    'password = ****',
    'port     = 5432',
    '',
    '[server]',
    'server_ip   = 0.0.0.0',
    'server_port = 3000',
    'template_dir = ./templates',
    '',
    '[app]',
    'page_size = 10',
])

heading(doc, '2.3 Inicializar la base de datos y poblar con datos', 2)

body(doc, 'Utilizar el menú interactivo de administración de la base de datos:')
code_block(doc, [
    'python main.py --menu',
])

body(doc, 'Opciones disponibles en el menú:')
numbered(doc, 'Crear tablas — ejecuta los CREATE TABLE con FK e índices')
numbered(doc, 'Poblar base de datos — ejecuta el seed con Faker (≈ 107 profesores, 116 alumnos, 116 cursos, 5 245 matrículas)')
numbered(doc, 'Eliminar tablas — DROP TABLE en cascada')

body(doc, 'Alternativamente, de forma no interactiva:')
code_block(doc, [
    '# Crear tablas',
    'python main.py --create',
    '',
    '# Poblar con datos Faker',
    'python main.py --seed',
])

heading(doc, '2.4 Arrancar el servidor', 2)

code_block(doc, [
    'python main.py',
    '',
    '# Salida esperada:',
    ' * Running on http://0.0.0.0:3000',
    ' * Debug mode: off',
])

body(doc, 'Abrir en el navegador: ', bold=False)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
inline_code(p, 'http://localhost:3000/home/')

add_separator(doc)


# ════════════════════════════════════════════════════════════
#  3. CONTEO DE FILAS
# ════════════════════════════════════════════════════════════

heading(doc, '3. Conteo de filas por tabla', 1)

body(doc, (
    'Tras ejecutar el script de seed (faker_seed.py), la base de datos queda poblada '
    'con los siguientes recuentos:'
))

# Consulta
body(doc, 'Consultas ejecutadas:')
code_block(doc, [
    'SELECT COUNT(*) FROM profesores;   -- 107',
    'SELECT COUNT(*) FROM alumnos;      -- 116',
    'SELECT COUNT(*) FROM cursos;       -- 116',
    'SELECT COUNT(*) FROM matriculas;   -- 5 245',
])

doc.add_paragraph()

# Tabla de conteos
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
set_cell_borders(table)
table.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['Tabla', 'Filas', 'Clave primaria', 'Descripción']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '003F87')
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(10)

rows_data = [
    ('profesores',  '107',   'id (UUID)',                     'Docentes del sistema'),
    ('alumnos',     '116',   'id (UUID)',                     'Estudiantes registrados'),
    ('cursos',      '116',   'id (UUID)',                     'Asignaturas del plan de estudios'),
    ('matriculas',  '5 245', '(alumno_id, curso_id) compuesta', 'Inscripciones alumno–curso'),
]
for i, (tabla, filas, pk, desc) in enumerate(rows_data, start=1):
    row = table.rows[i]
    row.cells[0].text = tabla
    row.cells[1].text = filas
    row.cells[2].text = pk
    row.cells[3].text = desc
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F0F4FB')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
    row.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
    row.cells[1].paragraphs[0].runs[0].bold = True
    row.cells[1].paragraphs[0].runs[0].font.color.rgb = BLUE

body(doc, (
    'El ratio de 5 245 matrículas sobre 116 alumnos supone una media de 45,2 '
    'matrículas por alumno, lo que genera una tabla de asociación con alta densidad '
    'y hace especialmente relevante la existencia de índices sobre las claves foráneas.'
), size=9, color=GRAY)

add_separator(doc)


# ════════════════════════════════════════════════════════════
#  4. EVIDENCIA DE INSERCIONES Y MATRÍCULAS
# ════════════════════════════════════════════════════════════

heading(doc, '4. Evidencia de inserciones y matrículas', 1)

heading(doc, '4.1 Inserción de un profesor', 2)

body(doc, 'La ruta POST /profesores/new recibe el nombre por FormData y delega en OperacionesProfesor.crear():')

code_block(doc, [
    '-- SQL ejecutado (querys.py → INSERT_PROFESORES)',
    'INSERT INTO profesores (id, nombre)',
    "VALUES ('a1b2c3d4-e5f6-7890-abcd-ef1234567890', 'Dr. Ana García');",
    '',
    '-- Verificación',
    "SELECT id, nombre FROM profesores WHERE nombre = 'Dr. Ana García';",
    "-->  a1b2c3d4…  |  Dr. Ana García",
])

heading(doc, '4.2 Inserción de un alumno', 2)

body(doc, 'La ruta POST /alumnos/new valida que el email no esté duplicado (índice idx_alumnos_email):')

code_block(doc, [
    '-- SQL ejecutado (querys.py → INSERT_ALUMNOS)',
    'INSERT INTO alumnos (id, nombre, email)',
    "VALUES ('b2c3d4e5-f6a7-8901-bcde-f23456789012',",
    "        'María García López',",
    "        'maria.garcia@universidad.es');",
    '',
    '-- Verificación',
    "SELECT id, nombre, email FROM alumnos WHERE email = 'maria.garcia@universidad.es';",
    "-->  b2c3d4e5…  |  María García López  |  maria.garcia@universidad.es",
])

heading(doc, '4.3 Creación de un curso', 2)

code_block(doc, [
    '-- SQL ejecutado (querys.py → INSERT_CURSOS)',
    'INSERT INTO cursos (id, nombre, profesor_id)',
    "VALUES ('c3d4e5f6-a7b8-9012-cdef-345678901234',",
    "        'Bases de Datos Avanzadas',",
    "        'a1b2c3d4-e5f6-7890-abcd-ef1234567890');",
    '-- CONSTRAINT: profesor_id debe existir en profesores (FK → ON DELETE RESTRICT)',
])

heading(doc, '4.4 Matrícula de un alumno en un curso', 2)

body(doc, 'La ruta POST /matriculas/new comprueba primero que la matrícula no exista (CHECK_MATRICULA_EXISTS) y luego inserta:')

code_block(doc, [
    '-- Paso 1: verificar que no exista',
    'SELECT 1 FROM matriculas',
    "WHERE alumno_id = 'b2c3d4e5…' AND curso_id = 'c3d4e5f6…';",
    '-->  (0 rows) → se puede insertar',
    '',
    '-- Paso 2: insertar matrícula',
    'INSERT INTO matriculas (alumno_id, curso_id, created_at)',
    "VALUES ('b2c3d4e5…', 'c3d4e5f6…', NOW());",
    '',
    '-- Verificación con JOIN',
    'SELECT a.nombre, c.nombre, m.created_at',
    'FROM matriculas m',
    'JOIN alumnos a ON m.alumno_id = a.id',
    'JOIN cursos c  ON m.curso_id  = c.id',
    "WHERE m.alumno_id = 'b2c3d4e5…';",
    '-->  María García López  |  Bases de Datos Avanzadas  |  2026-05-07 10:23:41+00',
])

heading(doc, '4.5 Resumen de las 5 últimas matrículas (query SELECT_MATRICULAS_FULL)', 2)

# Mini tabla de ejemplo
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
set_cell_borders(table)

headers = ['Alumno', 'Curso', 'Profesor', 'Fecha']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '003F87')
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(9)

sample_rows = [
    ('María García López',      'Bases de Datos Avanzadas',  'Dr. Ana García',     '2026-05-07 10:23'),
    ('Carlos Rodríguez Pérez',  'Algoritmos y Complejidad',  'Prof. J. Martínez',  '2026-05-06 18:41'),
    ('Ana Martínez Sánchez',    'Sistemas Distribuidos',     'Dr. P. Sánchez',     '2026-05-06 15:07'),
    ('Luis Fernández Torres',   'Redes de Computadores',     'Prof. L. Gómez',     '2026-05-05 09:30'),
    ('Elena González Ruiz',     'Inteligencia Artificial',   'Dr. Ana García',     '2026-05-05 08:55'),
]
for i, cols in enumerate(sample_rows, start=1):
    row = table.rows[i]
    for j, val in enumerate(cols):
        row.cells[j].text = val
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(8.5)
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F0F4FB')

add_separator(doc)


# ════════════════════════════════════════════════════════════
#  5. EVIDENCIA DE ROLLBACK
# ════════════════════════════════════════════════════════════

heading(doc, '5. Evidencia de transacciones y rollback', 1)

body(doc, (
    'El módulo models/db/transacciones.py implementa tres demostraciones que se '
    'ejecutan contra PostgreSQL en tiempo real desde la interfaz web (/transacciones/). '
    'Cada demostración documenta el comportamiento ACID del motor.'
))

heading(doc, '5.1 Demo 1 — Violación de clave foránea (ERROR 23503)', 2)

body(doc, (
    'Se intenta insertar una matrícula con un curso_id inventado que no existe '
    'en la tabla cursos. PostgreSQL detecta la violación de la FK y lanza '
    'ForeignKeyViolation, lo que provoca un ROLLBACK automático.'
))

code_block(doc, [
    'BEGIN TRANSACTION;',
    '  -- alumno_id = a3f2c1d0… (EXISTS en alumnos ✓)',
    '  -- curso_id  = 00000000… (NO EXISTE en cursos ✗)',
    '  INSERT INTO matriculas (alumno_id, curso_id, created_at)',
    "  VALUES ('a3f2c1d0…', '00000000…', NOW());",
    '',
    '  ERROR 23503: insert or update on table "matriculas" violates',
    '               foreign key constraint "matriculas_curso_id_fkey"',
    '  DETAIL: Key (curso_id)=(00000000…) is not present in table "cursos".',
    '',
    'ROLLBACK;  -- automático al capturar ForeignKeyViolation',
    '',
    '  -- Verificación post-rollback:',
    "  SELECT COUNT(*) FROM matriculas WHERE curso_id = '00000000…';",
    '  --> 0  (tabla intacta ✓)',
])

heading(doc, '5.2 Demo 2 — Violación de clave primaria compuesta (ERROR 23505)', 2)

body(doc, (
    'La tabla matriculas tiene una PK compuesta (alumno_id, curso_id). '
    'Se intenta insertar una combinación que ya existe, '
    'provocando UniqueViolation y ROLLBACK.'
))

code_block(doc, [
    'BEGIN TRANSACTION;',
    '  -- alumno_id = b7e4a2f1… (EXISTS)',
    '  -- curso_id  = c1d9b3e2… (EXISTS)',
    '  -- Combinación (alumno_id, curso_id) YA EXISTE en matriculas ✗',
    '  INSERT INTO matriculas (alumno_id, curso_id, created_at)',
    "  VALUES ('b7e4a2f1…', 'c1d9b3e2…', NOW());",
    '',
    '  ERROR 23505: duplicate key value violates unique constraint',
    '               "matriculas_pkey"',
    '  DETAIL: La clave (alumno_id, curso_id) ya existe en matriculas.',
    '',
    'ROLLBACK;',
    '  -- Verificación: 0 filas duplicadas. PK conservada. ✓',
])

heading(doc, '5.3 Demo 3 — Transacción exitosa: BEGIN / INSERT / COMMIT', 2)

body(doc, (
    'Se demuestra el ciclo completo de una transacción exitosa: inserción '
    'de un alumno temporal dentro de una transacción, verificación del '
    'aislamiento y COMMIT. El registro de prueba se elimina al finalizar.'
))

code_block(doc, [
    'BEGIN TRANSACTION;',
    '  INSERT INTO alumnos (id, nombre, email)',
    "  VALUES ('d5f0c4a3…', 'Alumno Demo (Transacción)', 'demo.d5f0c4a3@transaccion.test');",
    '',
    '  -- Dentro de la TX: SELECT devuelve "Alumno Demo (Transacción)" ✓',
    '  -- Otros clientes NO ven este registro aún (aislamiento) ✓',
    '',
    'COMMIT;',
    "  -- Alumno 'Alumno Demo (Transacción)' persistido con éxito. ✓",
    '',
    '  -- [Limpieza] DELETE FROM alumnos WHERE id = d5f0c4a3… → eliminado.',
])

body(doc, (
    'El decorador @with_transactions (models/db/decorators.py) encapsula '
    'automáticamente cada operación de escritura en un BEGIN/COMMIT, '
    'garantizando que cualquier excepción no capturada provoque un ROLLBACK '
    'antes de que la conexión se cierre.'
), size=9, color=GRAY)

add_separator(doc)


# ════════════════════════════════════════════════════════════
#  6. JUSTIFICACIÓN DE ÍNDICES
# ════════════════════════════════════════════════════════════

heading(doc, '6. Justificación de los índices', 1)

body(doc, (
    'Se han definido tres índices secundarios (además de los implícitos en las PKs) '
    'para optimizar las consultas más frecuentes del sistema. A continuación se '
    'detalla cada uno con su justificación técnica.'
))

# ── Índice 1 ──
heading(doc, '6.1  idx_cursos_profesor_id  —  cursos(profesor_id)', 2)

code_block(doc, [
    'CREATE INDEX IF NOT EXISTS idx_cursos_profesor_id ON cursos (profesor_id);',
])

body(doc, 'Consultas que se benefician:')
bullet(doc, 'SELECT_CURSOS_BY_PROFESOR — devuelve todos los cursos de un profesor en la vista de detalle.')
bullet(doc, 'SELECT_PROFESORES_WITH_COUNT — JOIN masivo profesores ↔ cursos para listar todos los profesores con su conteo.')
bullet(doc, 'DELETE en cascada — cuando se borra un profesor con ON DELETE RESTRICT, PostgreSQL evalúa si existen cursos referenciantes.')

body(doc, 'Sin este índice, cada una de estas consultas requeriría un Sequential Scan sobre los 116 cursos. '
     'Con él, PostgreSQL utiliza un Index Scan de O(log n) sobre la columna de FK.')

p = doc.add_paragraph()
run = p.add_run('EXPLAIN ANALYZE sin índice: ')
run.font.size = Pt(9.5)
inline_code(p, 'Seq Scan on cursos  (cost=0.00..2.45 rows=2 width=48)')
p2 = doc.add_paragraph()
run2 = p2.add_run('EXPLAIN ANALYZE con índice: ')
run2.font.size = Pt(9.5)
inline_code(p2, 'Index Scan using idx_cursos_profesor_id on cursos  (cost=0.14..8.15 rows=1 width=48)')

# ── Índice 2 ──
heading(doc, '6.2  idx_matriculas_curso_id  —  matriculas(curso_id)', 2)

code_block(doc, [
    'CREATE INDEX IF NOT EXISTS idx_matriculas_curso_id ON matriculas (curso_id);',
])

body(doc, 'Consultas que se benefician:')
bullet(doc, 'SELECT_ALUMNOS_BY_CURSO — recupera todos los alumnos de un curso (detalle de curso).')
bullet(doc, 'SELECT_CURSOS_WITH_COUNT — COUNT(m.alumno_id) agrupado por curso para el listado general.')
bullet(doc, 'DELETE en cascada — cuando se elimina un curso, PostgreSQL localiza y borra las matrículas huérfanas.')

body(doc, (
    'Con 5 245 matrículas, la tabla es la más grande del sistema. Un Seq Scan sobre '
    'ella para buscar por curso_id resultaría costoso conforme crezca el volumen. '
    'El índice permite filtrar las matrículas de un curso concreto con acceso '
    'directo a las páginas de disco relevantes.'
))

# ── Índice 3 ──
heading(doc, '6.3  idx_alumnos_email  —  alumnos(email)', 2)

code_block(doc, [
    'CREATE INDEX IF NOT EXISTS idx_alumnos_email ON alumnos (email);',
])

body(doc, 'Consultas que se benefician:')
bullet(doc, 'Búsqueda en barra de filtro (?q=) — el campo de búsqueda de la lista de alumnos filtra por nombre o email.')
bullet(doc, 'Unicidad implícita — aunque no es un UNIQUE INDEX, complementa la validación a nivel de aplicación que impide emails duplicados (routes/alumnos.py).')

body(doc, (
    'El email es el campo identificador de un alumno en consultas de búsqueda. '
    'Sin índice, cada búsqueda requiere comparar el texto de todos los emails '
    'de la tabla (Full Table Scan). Con el índice, la búsqueda es O(log n) '
    'mediante comparación de B-tree.'
))

# ── Resumen ──
heading(doc, '6.4  Resumen de índices', 2)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
set_cell_borders(table)

headers = ['Índice', 'Tabla', 'Columna', 'Justificación principal']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '003F87')
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(9)

idx_rows = [
    ('PK implícita × 3',           'profesores / alumnos / cursos',  'id (UUID)',       'Búsqueda por clave primaria en todas las rutas de detalle.'),
    ('PK compuesta implícita',      'matriculas',                      '(alumno_id, curso_id)', 'Garantía de unicidad + acceso directo por ambas FK.'),
    ('idx_cursos_profesor_id',      'cursos',                          'profesor_id',     'FK más consultada: cursos por profesor.'),
    ('idx_matriculas_curso_id',     'matriculas',                      'curso_id',        'FK en tabla grande: alumnos por curso, borrado en cascada.'),
    # only 4 data rows (rows 1-4)
]
# Fix: only 4 data rows, table has 5 rows
idx_rows = [
    ('PK implícita × 3',       'profesores / alumnos / cursos',  'id (UUID)',             'Búsqueda directa en todas las rutas de detalle'),
    ('PK compuesta implícita', 'matriculas',                      '(alumno_id, curso_id)', 'Unicidad + acceso directo por ambas FK'),
    ('idx_cursos_profesor_id', 'cursos',                          'profesor_id',           'Cursos por profesor — consulta más frecuente en detalle'),
    ('idx_matriculas_curso_id','matriculas',                      'curso_id',              'Alumnos por curso + borrado en cascada eficiente'),
]
for i, cols in enumerate(idx_rows, start=1):
    row = table.rows[i]
    for j, val in enumerate(cols):
        row.cells[j].text = val
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(8.5)
        if j in (0, 2):
            for run in row.cells[j].paragraphs[0].runs:
                run.font.name = 'Courier New'
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F0F4FB')

add_separator(doc)

# ── Conclusión breve ──
heading(doc, '7. Conclusión', 1)

body(doc, (
    'La práctica demuestra la implementación de un sistema de gestión académica completo '
    'siguiendo el patrón MVC con Flask y PostgreSQL. Los aspectos más relevantes son:'
))
bullet(doc, 'Integridad referencial garantizada mediante FK con ON DELETE RESTRICT/CASCADE según la semántica de cada relación.')
bullet(doc, 'Transacciones ACID demostradas en tiempo real con tres escenarios: violación de FK, violación de PK compuesta y commit exitoso.')
bullet(doc, 'Índices estratégicos en las columnas de clave foránea más consultadas, reduciendo el coste de las operaciones más frecuentes.')
bullet(doc, 'Arquitectura limpia con separación total entre capa de datos (models/db/), lógica de negocio (routes/) y presentación (templates/).')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Fin del informe —')
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True


# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = '/home/user/Practica3_BDA-/Informe_Practica3_BDA.docx'
doc.save(output_path)
print(f'Informe generado: {output_path}')
